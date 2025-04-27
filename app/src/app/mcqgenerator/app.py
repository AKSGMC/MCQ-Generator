"""
A App for Students and Teachers to generate Multiple Choice Questions
"""

import toga
from toga.style import Pack
from toga.style.pack import COLUMN, CENTER,ROW
from google import genai #for generating questions and answers
from fpdf import FPDF #for pdf Generation
from docx import Document #for ms word 
import easygui # for selecting files
import platform
import os

client = genai.Client(api_key='AIzaSyDFhA1whgmV9FRuGDlj-Q3_hfIApApGNmA')

class MCQGenerator(toga.App):
    def startup(self):
        self.main_box = toga.Box(style=Pack(direction=COLUMN, alignment=CENTER))
        
        # Dropdown box for selecting modes
        mode_label = toga.Label("Select Mode", style=Pack(text_align=CENTER, font_size=20, padding_bottom=20))
        self.dropdown = toga.Selection(items=["Select Mode", "Hint generator", "Paragraph Mode"], style=Pack(width=200, padding_bottom=20))
        self.dropdown.on_change = self.on_select_mode
        
        # adding elements to main box
        self.main_box.add(mode_label)
        self.main_box.add(self.dropdown)
        
        # setting main window
        self.main_window = toga.MainWindow(title=self.formal_name)
        self.main_window.content = self.main_box
        self.main_window.show()

    def on_select_mode(self, widget):
        self.selected_mode = widget.value
        self.main_box.remove(*self.main_box.children[2:]) 
        
        if self.selected_mode == "PDF Mode":
            self.show_pdf_mode_content(widget)
        elif self.selected_mode == "Hint generator":
            self.show_hint_generator_content(widget)
        elif self.selected_mode == "Paragraph Mode":
            self.show_paragraph_mode_content(widget)
        elif self.selected_mode == "Select Mode":
            select_mode_label = toga.Label("Select Mode", style=Pack(text_align=CENTER, font_size=15, padding_bottom=20))
            select_mode_Text = toga.Label("Please Select a mode from above dropdown box", style=Pack(text_align=CENTER, font_size=10, padding_bottom=30))
            self.main_box.add(select_mode_label)
            self.main_box.add(select_mode_Text)

     #output filetype page       
    def output_file_type_page(self,widget):
        Output_file_label=toga.Label("Select file type: ",style=Pack(font_size=15,padding_top=15))
        self.select_file_type=toga.Selection(items=["--Select--","PDF TYPE","TXT FILE TYPE","MS Word TYPE"],style=Pack(padding_top=20,width=150))
        generate_button=toga.Button("Generate!",on_press=self.check_mode,style=Pack(padding_top=15,width=100))
        #creating a new box
        output_file_page= toga.Box(style=Pack(direction=ROW,alignment=CENTER))
        #adding to new box
        output_file_page.add(Output_file_label)
        output_file_page.add(self.select_file_type)
        #adding new box to main box
        self.main_box.add(output_file_page)
        #adding to main box
        self.main_box.add(generate_button)
    
    
    def check_mode(self,widget):
        self.main_box.remove(*self.main_box.children[1:])
        #checking the selected mode to call different generate function
        if self.selected_mode=='Paragraph Mode':
            self.Question_Generate(widget)
        if self.selected_mode=='PDF Mode':
            self.PDF_Generate_Answers(widget)
        if self.selected_mode=='Hint generator':
            self.Hint_Generate(widget)

    #code for selecting file type
    def file_type_on_change(self,widget):
        self.file_selected_mode= self.select_file_type.value
        if self.file_selected_mode=='PDF TYPE':
            self.pdf_file_question_write(widget)
        if self.file_selected_mode=='TXT FILE TYPE':
            self.txt_file_question_write(widget)
        if self.file_selected_mode=='MS Word TYPE':
            self.ms_word_file_question_write(widget)
        

    #code for pdf mode
    def show_pdf_mode_content(self,widget):
        pdf_label = toga.Label("PDF Mode - Generates answer only", style=Pack(text_align=CENTER, font_size=15, padding_bottom=20))
        pdf_select=toga.Button("Select PDF",style=Pack(padding_bottom=10,width=100))
        pdf_select.on_press=self.PDF_File_select
        self.main_box.add(pdf_select)
        self.main_box.add(pdf_label)

    def PDF_File_select(self,widget):
        pdf_selected= easygui.fileopenbox(msg='Select PDF file to generate answers',multiple=False,filetypes='pdf')
        self.pdf= client.files.upload(file=pdf_selected)
        Generate_button = toga.Button("Generate!",on_press=self.call_generate,style=Pack(padding_top=15,width=100))
        # add to main box
        self.main_box.add(Generate_button)
    def call_generate(self,widget):
        self.PDF_on_submit(widget)

    #code for hint generator mode
    def show_hint_generator_content(self, widget):
        hint_label = toga.Label("Hint Generator", style=Pack(text_align=CENTER, font_size=15, padding_bottom=20))
        hint_input= toga.MultilineTextInput(style=Pack(width=500,height=200,padding_top=20))
        hint_button_submit=toga.Button("Submit",style=Pack(padding_top=20,width=100))
        self.hint_paragraph=hint_input
        hint_button_submit.on_press=self.hint_generator_on_submit
        hint_input.placeholder="Enter Paragraph Here to convert into hints"

        #adding to the mainbox
        self.main_box.add(hint_label)
        self.main_box.add(hint_input)
        self.main_box.add(hint_button_submit)

    #code for paragraph mode
    def show_paragraph_mode_content(self,widget):
        paragraph_label = toga.Label("Paragraph Mode", style=Pack(text_align=CENTER, font_size=15, padding_bottom=20,padding_top=10))
        paragraph_input = toga.MultilineTextInput(style=Pack(width=500, height=200, padding_top=20))
        paragraph_button_submit= toga.Button("Submit",style=Pack(padding_top=20,width=100))
        self.paragraph=paragraph_input
        paragraph_button_submit.on_press= self.paragraph_on_submit
        paragraph_input.placeholder="Enter the paragraph here"

        #Add to the mainbox 
        self.main_box.add(paragraph_label)        
        self.main_box.add(paragraph_input)
        self.main_box.add(paragraph_button_submit)

    # action to done on clicking submit buttons
    def hint_generator_on_submit(self,widget):
        self.main_box.remove(*self.main_box.children)
        self.hint_paragraph_value=self.hint_paragraph.value
        hint_label = toga.Label("Hint Generator", style=Pack(text_align=CENTER, font_size=15, padding_bottom=20))
        
        #adding to the box
        self.main_box.add(hint_label)
        #calling output file paga
        self.output_file_type_page(widget)

    def PDF_on_submit(self,widget):
        self.main_box.remove(*self.main_box.children)
        self.output_file_type_page(widget)

    def paragraph_on_submit(self,widget):
        self.main_box.remove(*self.main_box.children)
        self.paragraph_value=self.paragraph.value
        paragraph_label = toga.Label("Paragraph Mode", style=Pack(text_align=CENTER, font_size=20, padding_bottom=15,padding_top=10))
        paragraph_question_number_label = toga.Label("Enter the number of questions you want to generate", style=Pack(text_align=CENTER, font_size=15, padding_top=15,padding_bottom=15))
        paragraph_choice_number_label = toga.Label("Enter the number of choices you want to generate", style=Pack(text_align=CENTER, font_size=15, padding_bottom=15,padding_top=15))
        paragraph_input_question_number = toga.TextInput(style=Pack(width=500, padding_bottom=15))
        paragraph_input_choice_number = toga.TextInput(style=Pack(width=500, padding_bottom=15))
        paragraph_button_submit_page2= toga.Button("Submit",style=Pack(width=100))
        self.question_value=paragraph_input_question_number
        self.choice_value=paragraph_input_choice_number
        paragraph_button_submit_page2.on_press= self.paragraph_on_submit_page2
        
        #adding to the mainbox
        self.main_box.add(paragraph_label)
        self.main_box.add(paragraph_question_number_label)
        self.main_box.add(paragraph_input_question_number)
        self.main_box.add(paragraph_choice_number_label)
        self.main_box.add(paragraph_input_choice_number)
        self.main_box.add(paragraph_button_submit_page2)

    def paragraph_on_submit_page2(self,widget):
            self.main_box.remove(*self.main_box.children[1:])
            self.paragraph_question_value=self.question_value.value
            self.paragraph_choice_value=self.choice_value.value
            self.output_file_type_page(widget)

    def Question_Generate(self,widget):
        #generating questions
        try:
            Generated_Question= client.models.generate_content(
                model="gemini-2.0-flash",
                contents='''Generate {} MCQ Question with {} choices with one correct choice
                    and don't 
                    highlight the answer and question. And also generate the answer below each question with explanation
                 from the paragraph: {}'''.format(self.paragraph_question_value,self.paragraph_choice_value,self.paragraph_value),
            )
            self.Generated=Generated_Question.text
        except Exception as e:
            wait_label=toga.Label("Error in generating answer.\nPlease Check your Internet Connection and reopen the app\n\n If you think this is an issue please report it on github ",style=Pack(font_size=15,padding_top=10,text_align=CENTER))
            self.main_box.add(wait_label)
            print(str(e))
            
        #calling file type caller
        self.file_type_on_change(widget)

    def Hint_Generate(self,widget):
        try:
            generate_hint= client.models.generate_content(
                model='gemini-2.0-flash',
                contents='''Generate hints from the given paragraph: {}'''.format(self.hint_paragraph_value)
            )
            self.Generated= generate_hint.text
        except Exception as e:
            wait_label=toga.Label("Error in generating answer.\nPlease Check your Internet Connection and reopen the app\n\n If you think this is an issue please report it on github ",style=Pack(font_size=15,padding_top=10,text_align=CENTER))
            self.main_box.add(wait_label)
            print(str(e))

        self.file_type_on_change(widget)

    def PDF_Generate_Answers(self,widget):
        try:
            self.main_box.remove(*self.main_box.children)
            generate_pdf_answers= client.models.generate_content(
                model='gemini-2.0-flash',
                contents='''Generate Answers from given content with explanation: {}'''.format(self.pdf)
            )
            self.Generated= generate_pdf_answers.text
        except Exception as e:
            wait_label=toga.Label("Error in generating answer.\nPlease Check your Internet Connection and reopen the app\n\n If you think this is an issue please report it on github ",style=Pack(font_size=15,padding_top=10,text_align=CENTER))
            self.main_box.add(wait_label)
            print(str(e))

        self.file_type_on_change(widget)

    #file creation

    # writting pdf file
    def pdf_file_question_write(self,widget):
            os_name = platform.system()
            if os_name == "Windows":
                # Use the Desktop directory
                download_path = os.path.join(os.environ['USERPROFILE'],'Downloads','Generated_Questions.pdf')
                output_path=download_path
            elif os_name in ["Linux", "Darwin"]:  # macOS is "Darwin"
                desktop_path = os.path.join(os.path.expanduser("~"), 'Desktop', 'Generated_Questions.pdf')
                output_path=desktop_path
            elif "ANDROID_BOOTLOGO" in os.environ:
                # If running on Android, use the Downloads folder
                download_path = "/storage/emulated/0/Download/Generated_Questions.pdf"
                output_path=download_path
            else:
                raise Exception("Unsupported operating system")
            #Display Information
            wait_label=toga.Label(f"Generated Question saved at location \n {output_path}",style=Pack(font_size=15,padding_top=10,text_align=CENTER))
            self.main_box.add(wait_label)
            # Create a PDF using fpdf2
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font('Helvetica',size=12)
            txt_Generated=self.Generated
            pdf.write(text=txt_Generated)
            pdf.set_author("MCQ Generator")
            pdf.set_title("Generated Questions")
            # Save the PDF
            pdf.output(output_path)

    # Writting text file
    def txt_file_question_write(self,widget):
        os_name = platform.system()
        if os_name == "Windows":
            # Use the Desktop directory
            download_path = os.path.join(os.environ['USERPROFILE'],'Downloads','Generated.txt')
            output_path=download_path
        elif os_name in ["Linux", "Darwin"]:  # macOS is "Darwin"
            desktop_path = os.path.join(os.path.expanduser("~"), 'Desktop', 'Generated.txt')
            output_path=desktop_path
        elif "ANDROID_BOOTLOGO" in os.environ:
          # If running on Android, use the Downloads folder
            download_path = "/storage/emulated/0/Download/Generated.txt"
            output_path=download_path
        else:
          raise Exception("Unsupported operating system")
        #Display Information
        wait_label=toga.Label(f"Generated Question saved at location \n {output_path}",style=Pack(font_size=15,padding_top=10,text_align=CENTER))
        self.main_box.add(wait_label)

        #file writting
        with open(output_path,'w') as txt:
            txt.write(self.Generated)

    def ms_word_file_question_write(self,widget):
        os_name = platform.system()
        if os_name == "Windows":
            # Use the Desktop directory
            download_path = os.path.join(os.environ['USERPROFILE'],'Downloads','Generated.docx')
            output_path=download_path
        elif os_name in ["Linux", "Darwin"]:  # macOS is "Darwin"
            desktop_path = os.path.join(os.path.expanduser("~"), 'Desktop', 'Generated.docx')
            output_path=desktop_path
        elif "ANDROID_BOOTLOGO" in os.environ:
            # If running on Android, use the Downloads folder
            download_path = "/storage/emulated/0/Download/Generated.docx"
            output_path=download_path
        else:
            raise Exception("Unsupported operating system")
        #Display Information
        wait_label=toga.Label(f"Generated Question saved at location \n {output_path}",style=Pack(font_size=15,padding_top=10,text_align=CENTER))
        self.main_box.add(wait_label)
        # Document handle
        doc = Document()
        doc.add_heading("Generated Content")
        doc.add_paragraph(self.Generated)
        doc.save(output_path)

def main():
    return MCQGenerator()
